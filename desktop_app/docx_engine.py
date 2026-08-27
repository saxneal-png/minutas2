import os
from typing import Dict, Any, List
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex: str):
    """Set background color of a table cell in docx."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner cell padding/margins."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

class DocxEngine:
    def __init__(self, templates_dir: str = None):
        self.templates_dir = templates_dir or os.path.join(os.path.dirname(__file__), "templates")
        os.makedirs(self.templates_dir, exist_ok=True)
        self.default_template_path = os.path.join(self.templates_dir, "plantilla_oficial_doh.docx")
        self.ensure_default_template()

    def ensure_default_template(self):
        """Generates a high-quality professional default DOCX template if missing."""
        if os.path.exists(self.default_template_path):
            return

        doc = docx.Document()
        
        # Set margins to 2 cm
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.9)
            section.right_margin = Inches(0.9)

        # Main Title Header
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_after = Pt(2)
        run_org = title_p.add_run("DIRECCIÓN DE OBRAS HIDRÁULICAS - MOP\n")
        run_org.font.name = "Arial"
        run_org.font.size = Pt(11)
        run_org.font.bold = True
        run_org.font.color.rgb = RGBColor(30, 58, 138) # Navy Blue

        run_sub = title_p.add_run("SERVICIO DE APOYO EXPERTO Y COORDINACIÓN TERRITORIAL\n")
        run_sub.font.name = "Arial"
        run_sub.font.size = Pt(10)
        run_sub.font.bold = True
        run_sub.font.color.rgb = RGBColor(71, 85, 105)

        run_main = title_p.add_run("MINUTA DE REUNIÓN Y ACUERDOS TÉCNICOS")
        run_main.font.name = "Arial"
        run_main.font.size = Pt(14)
        run_main.font.bold = True
        run_main.font.color.rgb = RGBColor(15, 23, 42)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # Info Box Table
        table = doc.add_table(rows=6, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        fields = [
            ("PROYECTO / ASUNTO:", "{ASUNTO}"),
            ("FECHA:", "{FECHA}"),
            ("HORA:", "{HORA}"),
            ("LUGAR:", "{LUGAR}"),
            ("COORDINADOR(A):", "{COORDINADOR}"),
            ("ASISTENTES:", "{ASISTENTES}")
        ]

        for i, (label, val) in enumerate(fields):
            # Left cell (Label)
            c0 = table.cell(i, 0)
            c0.width = Inches(2.0)
            set_cell_background(c0, "1E3A8A") # Dark blue
            p0 = c0.paragraphs[0]
            p0.paragraph_format.space_after = Pt(2)
            p0.paragraph_format.space_before = Pt(2)
            r0 = p0.add_run(label)
            r0.font.name = "Arial"
            r0.font.size = Pt(9.5)
            r0.font.bold = True
            r0.font.color.rgb = RGBColor(255, 255, 255)

            # Right cell (Value)
            c1 = table.cell(i, 1)
            c1.width = Inches(4.7)
            set_cell_background(c1, "F8FAFC") # Light gray
            p1 = c1.paragraphs[0]
            p1.paragraph_format.space_after = Pt(2)
            p1.paragraph_format.space_before = Pt(2)
            r1 = p1.add_run(val)
            r1.font.name = "Arial"
            r1.font.size = Pt(9.5)
            r1.font.color.rgb = RGBColor(15, 23, 42)

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

        # Section 1: Context and Details
        h1 = doc.add_paragraph()
        h1.paragraph_format.space_before = Pt(12)
        h1.paragraph_format.space_after = Pt(4)
        r_h1 = h1.add_run("1. ANTECEDENTES Y TEMAS TRATADOS")
        r_h1.font.name = "Arial"
        r_h1.font.size = Pt(11)
        r_h1.font.bold = True
        r_h1.font.color.rgb = RGBColor(30, 58, 138)

        p_det = doc.add_paragraph()
        p_det.paragraph_format.space_after = Pt(12)
        p_det.paragraph_format.line_spacing = 1.15
        r_det = p_det.add_run("{DETALLES}")
        r_det.font.name = "Arial"
        r_det.font.size = Pt(10)

        # Section 2: Agreements Table
        h2 = doc.add_paragraph()
        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after = Pt(4)
        r_h2 = h2.add_run("2. MATRIZ DE ACUERDOS Y COMPROMISOS")
        r_h2.font.name = "Arial"
        r_h2.font.size = Pt(11)
        r_h2.font.bold = True
        r_h2.font.color.rgb = RGBColor(30, 58, 138)

        p_marker = doc.add_paragraph()
        r_mark = p_marker.add_run("{TABLA_COMPROMISOS}")
        r_mark.font.name = "Arial"
        r_mark.font.size = Pt(9)
        r_mark.font.italic = True

        doc.save(self.default_template_path)

    def generate_docx(self, data: Dict[str, Any], output_path: str, template_path: str = None) -> str:
        """
        Renders minute data into a formatted Word document.
        Uses the provided template or the built-in official template.
        """
        tpl_path = template_path or self.default_template_path
        if not os.path.exists(tpl_path):
            self.ensure_default_template()
            tpl_path = self.default_template_path

        doc = docx.Document(tpl_path)

        # Prepare replacement mapping with flexible aliases and double-bracket support
        detalles_val = str(data.get("detalles", "Sin antecedentes registrados."))
        asunto_val = str(data.get("asunto", "No especificado"))
        fecha_val = str(data.get("fecha", "No especificado"))
        hora_val = str(data.get("hora", "No especificado"))
        lugar_val = str(data.get("lugar", "No especificado"))
        coord_val = str(data.get("coordinador", "Coordinadora de Oficina de Actividades Territoriales"))
        asist_val = str(data.get("asistentes", "No especificado"))

        replacements = {
            # Single braces
            "{ASUNTO}": asunto_val,
            "{asunto}": asunto_val,
            "{FECHA}": fecha_val,
            "{fecha}": fecha_val,
            "{HORA}": hora_val,
            "{hora}": hora_val,
            "{LUGAR}": lugar_val,
            "{lugar}": lugar_val,
            "{COORDINADOR}": coord_val,
            "{coordinador}": coord_val,
            "{ASISTENTES}": asist_val,
            "{asistentes}": asist_val,
            "{DETALLES}": detalles_val,
            "{detalles}": detalles_val,
            "{DETALLES_REUNION}": detalles_val,
            "{DETALLES_DE_LA_REUNION}": detalles_val,
            "{ANTECEDENTES}": detalles_val,
            
            # Double braces
            "{{ASUNTO}}": asunto_val,
            "{{asunto}}": asunto_val,
            "{{FECHA}}": fecha_val,
            "{{fecha}}": fecha_val,
            "{{HORA}}": hora_val,
            "{{hora}}": hora_val,
            "{{LUGAR}}": lugar_val,
            "{{lugar}}": lugar_val,
            "{{COORDINADOR}}": coord_val,
            "{{coordinador}}": coord_val,
            "{{ASISTENTES}}": asist_val,
            "{{asistentes}}": asist_val,
            "{{DETALLES}}": detalles_val,
            "{{detalles}}": detalles_val,
            "{{DETALLES_REUNION}}": detalles_val,
            "{{DETALLES_DE_LA_REUNION}}": detalles_val,
            "{{ANTECEDENTES}}": detalles_val
        }

        # 1. Replace in all paragraphs
        for p in doc.paragraphs:
            for key, val in replacements.items():
                if key in p.text:
                    p.text = p.text.replace(key, val)

        filas = data.get("filas", [])
        table_placeholder_found = False

        # 2. Check and fill existing template tables if they contain row placeholders like {TEMA} / {COMPROMISO}
        for table in doc.tables:
            # First, check if any cell is a template row for compromisos
            row_to_template = None
            for row_idx, row in enumerate(table.rows):
                row_text = " ".join(c.text for c in row.cells)
                if "{TEMA}" in row_text or "{{TEMA}}" in row_text or "{COMPROMISO}" in row_text or "{{COMPROMISO}}" in row_text:
                    row_to_template = row_idx
                    table_placeholder_found = True
                    break

            if row_to_template is not None:
                # We found a table that has a row with {TEMA}, {COMPROMISO}, etc.
                tpl_row = table.rows[row_to_template]
                # If there are filas, we populate this row with the first fila and add new rows for remaining
                if filas:
                    for f_idx, fila in enumerate(filas):
                        target_row = tpl_row if f_idx == 0 else table.add_row()
                        # Populate cells
                        tema_text = str(fila.get("tema", ""))
                        comp_text = str(fila.get("compromiso", ""))
                        resp_text = str(fila.get("responsable", fila.get("plazo", "")))
                        plazo_text = str(fila.get("plazo", "No especificado"))
                        if "responsable" not in fila:
                            resp_text = "No especificado"

                        if f_idx == 0:
                            for cell in target_row.cells:
                                for p in cell.paragraphs:
                                    p.text = p.text.replace("{TEMA}", tema_text).replace("{{TEMA}}", tema_text)
                                    p.text = p.text.replace("{COMPROMISO}", comp_text).replace("{{COMPROMISO}}", comp_text)
                                    p.text = p.text.replace("{RESPONSABLE}", resp_text).replace("{{RESPONSABLE}}", resp_text)
                                    p.text = p.text.replace("{PLAZO}", plazo_text).replace("{{PLAZO}}", plazo_text)
                                    p.text = p.text.replace("{FECHA}", plazo_text).replace("{{FECHA}}", plazo_text)
                        else:
                            # Copy cell contents and styling from template row
                            for c_idx, cell in enumerate(target_row.cells):
                                if c_idx < len(tpl_row.cells):
                                    ref_text = tpl_row.cells[c_idx].text
                                    val = ""
                                    if "{TEMA}" in ref_text or "{{TEMA}}" in ref_text:
                                        val = tema_text
                                    elif "{COMPROMISO}" in ref_text or "{{COMPROMISO}}" in ref_text:
                                        val = comp_text
                                    elif "{RESPONSABLE}" in ref_text or "{{RESPONSABLE}}" in ref_text:
                                        val = resp_text
                                    elif "{PLAZO}" in ref_text or "{{PLAZO}}" in ref_text or "{FECHA}" in ref_text or "{{FECHA}}" in ref_text:
                                        val = plazo_text
                                    else:
                                        val = tema_text
                                    cell.paragraphs[0].text = val
                else:
                    for cell in tpl_row.cells:
                        for p in cell.paragraphs:
                            p.text = p.text.replace("{TEMA}", "Sin compromisos").replace("{{TEMA}}", "Sin compromisos")
                            p.text = p.text.replace("{COMPROMISO}", "-").replace("{{COMPROMISO}}", "-")
                            p.text = p.text.replace("{PLAZO}", "-").replace("{{PLAZO}}", "-")

            # Also do standard variable replacement for remaining cells in table
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, val in replacements.items():
                            if key in p.text:
                                p.text = p.text.replace(key, val)

        # 3. Handle Agreement Table insertion in paragraphs if {TABLA_COMPROMISOS} is used
        if not table_placeholder_found:
            for i, p in enumerate(doc.paragraphs):
                if "{TABLA_COMPROMISOS}" in p.text or "{{TABLA_COMPROMISOS}}" in p.text or "{filas}" in p.text or "{{filas}}" in p.text:
                    p.text = "" # Clear placeholder
                    table_placeholder_found = True
                    
                    # Create styled table
                    table = doc.add_table(rows=1 + len(filas), cols=5)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.autofit = False

                    # Header row
                    headers = ["N°", "Tema / Punto Tratado", "Acuerdo / Compromiso", "Responsable", "Plazo"]
                    widths = [Inches(0.4), Inches(1.8), Inches(2.5), Inches(1.2), Inches(0.8)]
                    
                    hdr_cells = table.rows[0].cells
                    for col_idx, (hdr_text, w) in enumerate(zip(headers, widths)):
                        cell = hdr_cells[col_idx]
                        cell.width = w
                        set_cell_background(cell, "1E3A8A")
                        set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
                        p_cell = cell.paragraphs[0]
                        p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [0, 4] else WD_ALIGN_PARAGRAPH.LEFT
                        r = p_cell.add_run(hdr_text)
                        r.font.name = "Arial"
                        r.font.size = Pt(9)
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)

                    # Data rows
                    for row_idx, fila in enumerate(filas):
                        row_cells = table.rows[row_idx + 1].cells
                        bg_color = "F1F5F9" if row_idx % 2 == 1 else "FFFFFF"
                        
                        values = [
                            str(row_idx + 1),
                            str(fila.get("tema", "")),
                            str(fila.get("compromiso", "")),
                            str(fila.get("responsable", fila.get("plazo", ""))),
                            str(fila.get("plazo", "No especificado"))
                        ]
                        
                        if "responsable" not in fila:
                            values[3] = "No especificado"
                            values[4] = str(fila.get("plazo", "No especificado"))

                        for col_idx, (val, w) in enumerate(zip(values, widths)):
                            cell = row_cells[col_idx]
                            cell.width = w
                            set_cell_background(cell, bg_color)
                            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
                            p_cell = cell.paragraphs[0]
                            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [0, 4] else WD_ALIGN_PARAGRAPH.LEFT
                            r = p_cell.add_run(val)
                            r.font.name = "Arial"
                            r.font.size = Pt(8.5)
                            r.font.color.rgb = RGBColor(15, 23, 42)

                    # Move table right after the paragraph marker
                    p._p.addnext(table._tbl)
                    break

        # If still no placeholder was found in custom template, append table at the end
        if not table_placeholder_found and filas:
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(12)
            r = h.add_run("Compromisos y Acuerdos:")
            r.font.bold = True

            table = doc.add_table(rows=1 + len(filas), cols=4)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            for idx, title in enumerate(["N°", "Tema", "Compromiso", "Plazo"]):
                hdr[idx].paragraphs[0].add_run(title).font.bold = True
            
            for row_idx, fila in enumerate(filas):
                r_cells = table.rows[row_idx + 1].cells
                r_cells[0].paragraphs[0].text = str(row_idx + 1)
                r_cells[1].paragraphs[0].text = str(fila.get("tema", ""))
                r_cells[2].paragraphs[0].text = str(fila.get("compromiso", ""))
                r_cells[3].paragraphs[0].text = str(fila.get("plazo", ""))

        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        doc.save(output_path)
        return output_path
