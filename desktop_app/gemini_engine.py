import os
import json
import re
import base64
from typing import Dict, Any, Tuple, Optional, List
import requests
import pypdf
import docx

class GeminiEngine:
    """
    Motor ultra-robusto de conexión con la API de Google Gemini.
    Soporta análisis individual y compilación unificada de MÚLTIPLES archivos
    (Word, PDF, imágenes, audios y notas) en una sola minuta consolidada.
    """
    def __init__(self, api_key: str = "", model_name: str = "gemini-2.0-flash", temperature: float = 0.1):
        self.api_key = api_key.strip()
        self.model_name = model_name or "gemini-2.0-flash"
        self.temperature = temperature
        self.timeout = 55 # Timeout generoso para compilaciones multi-archivo
        self.cached_models = []

    def set_api_key(self, api_key: str):
        self.api_key = api_key.strip()

    def get_available_models(self) -> Tuple[bool, List[str], str]:
        """Consulta en tiempo real los modelos disponibles habilitados para la API Key."""
        if not self.api_key:
            return False, [], "Ingresa tu API Key para descubrir los modelos disponibles."

        url = f"https://generativelanguage.googleapis.com/v1beta/models?key={self.api_key}"

        try:
            resp = requests.get(url, timeout=10)
            if resp.status_code == 200:
                data = resp.json()
                models_raw = data.get("models", [])
                valid_models = []

                for m in models_raw:
                    methods = m.get("supportedGenerationMethods", [])
                    if "generateContent" in methods:
                        name = m.get("name", "").replace("models/", "")
                        if "gemini" in name.lower() and not "vision" in name.lower() and not "embedding" in name.lower():
                            valid_models.append(name)

                priority_order = [
                    "gemini-2.0-flash",
                    "gemini-2.0-flash-exp",
                    "gemini-1.5-flash",
                    "gemini-1.5-pro",
                    "gemini-2.5-flash",
                    "gemini-3.7-flash"
                ]

                sorted_models = []
                for p in priority_order:
                    for vm in valid_models:
                        if vm == p and vm not in sorted_models:
                            sorted_models.append(vm)

                for vm in valid_models:
                    if vm not in sorted_models:
                        sorted_models.append(vm)

                if not sorted_models:
                    sorted_models = ["gemini-2.0-flash", "gemini-1.5-flash", "gemini-1.5-pro"]

                self.cached_models = sorted_models
                return True, sorted_models, f"Se detectaron {len(sorted_models)} modelos Gemini activos."
            else:
                err_data = resp.json().get("error", {})
                err_msg = err_data.get("message", resp.text)
                return False, ["gemini-2.0-flash", "gemini-1.5-flash"], f"Error consultando modelos: {err_msg}"
        except Exception as e:
            return False, ["gemini-2.0-flash", "gemini-1.5-flash"], f"Error de conexión: {str(e)}"

    def test_connection(self) -> Tuple[bool, str, List[str]]:
        if not self.api_key:
            return False, "Por favor ingresa una API Key de Gemini.", []

        ok_models, models_list, models_msg = self.get_available_models()
        if not ok_models:
            return False, models_msg, []

        test_model = self.model_name if self.model_name in models_list else models_list[0]
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{test_model}:generateContent?key={self.api_key}"
        payload = {
            "contents": [{
                "parts": [{"text": "Responde únicamente con la palabra 'OK'."}]
            }],
            "generationConfig": {
                "temperature": 0.0,
                "maxOutputTokens": 10
            }
        }

        try:
            resp = requests.post(url, json=payload, timeout=12)
            if resp.status_code == 200:
                return True, f"¡Conexión exitosa con {test_model}! ({len(models_list)} modelos listos)", models_list
            else:
                err_data = resp.json().get("error", {})
                err_msg = err_data.get("message", resp.text)
                return False, f"Error validando modelo {test_model}: {err_msg}", models_list
        except Exception as e:
            return False, f"Error al conectar con Google: {str(e)}", models_list

    def extract_text_from_file(self, file_path: str) -> Tuple[str, Optional[str], Optional[str]]:
        """Extrae texto o convierte archivos a base64."""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".docx":
            try:
                doc = docx.Document(file_path)
                full_text = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        full_text.append(p.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        row_vals = [c.text.strip() for c in row.cells if c.text.strip()]
                        if row_vals:
                            full_text.append(" | ".join(row_vals))
                extracted = "\n".join(full_text)
                if not extracted.strip():
                    raise Exception("El documento Word está vacío o no contiene texto legible.")
                return extracted, None, None
            except Exception as e:
                raise Exception(f"No se pudo leer el archivo Word ({os.path.basename(file_path)}): {e}")

        elif ext == ".pdf":
            try:
                reader = pypdf.PdfReader(file_path)
                pages_text = []
                for i, page in enumerate(reader.pages):
                    t = page.extract_text()
                    if t and t.strip():
                        pages_text.append(f"--- PÁGINA {i+1} ---\n" + t.strip())
                extracted = "\n\n".join(pages_text)
                if len(extracted.strip()) > 30:
                    return extracted, None, None
                
                with open(file_path, "rb") as f:
                    b64 = base64.b64encode(f.read()).decode("utf-8")
                return "", "application/pdf", b64
            except Exception:
                with open(file_path, "rb") as f:
                    b64 = base64.b64encode(f.read()).decode("utf-8")
                return "", "application/pdf", b64

        elif ext in [".txt", ".md", ".csv", ".log"]:
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read(), None, None
            except Exception as e:
                raise Exception(f"No se pudo leer el archivo de texto: {e}")

        elif ext in [".png", ".jpg", ".jpeg", ".webp", ".bmp"]:
            mime = "image/png" if ext == ".png" else "image/jpeg"
            with open(file_path, "rb") as f:
                b64 = base64.b64encode(f.read()).decode("utf-8")
            return "", mime, b64

        elif ext in [".mp3", ".wav", ".m4a", ".ogg", ".aac"]:
            mime_map = {
                ".mp3": "audio/mp3",
                ".wav": "audio/wav",
                ".m4a": "audio/m4a",
                ".ogg": "audio/ogg",
                ".aac": "audio/aac"
            }
            mime = mime_map.get(ext, "audio/mp3")
            with open(file_path, "rb") as f:
                b64 = base64.b64encode(f.read()).decode("utf-8")
            return "", mime, b64

        else:
            raise Exception(f"Formato no compatible: '{ext}'.")

    def _call_gemini_api(self, model_name: str, parts: List[Dict[str, Any]], system_prompt: str) -> str:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={self.api_key}"

        payload = {
            "contents": [{
                "parts": parts
            }],
            "generationConfig": {
                "temperature": self.temperature,
                "responseMimeType": "application/json"
            }
        }

        if system_prompt:
            payload["systemInstruction"] = {
                "parts": [{"text": system_prompt}]
            }

        response = requests.post(url, json=payload, timeout=self.timeout)
        
        if response.status_code == 200:
            data = response.json()
            candidates = data.get("candidates", [])
            if not candidates:
                feedback = data.get("promptFeedback", {})
                block_reason = feedback.get("blockReason", "Sin candidatos generados")
                raise Exception(f"La respuesta fue bloqueada por filtros de seguridad: {block_reason}")

            part = candidates[0].get("content", {}).get("parts", [{}])[0]
            raw_text = part.get("text", "")
            if not raw_text:
                raise Exception("Gemini devolvió una respuesta vacía.")
            return raw_text
        else:
            err_json = response.json().get("error", {})
            err_msg = err_json.get("message", response.text)
            err_code = err_json.get("code", response.status_code)
            raise Exception(f"Error {err_code} ({model_name}): {err_msg}")

    def analyze_compiled_sources(self, 
                                 file_paths: List[str], 
                                 raw_text: str, 
                                 system_prompt: str, 
                                 progress_callback=None) -> Dict[str, Any]:
        """
        Compila y sintetiza MÚLTIPLES archivos y/o notas directas en una única minuta completa.
        """
        if not self.api_key:
            raise Exception("No se ha configurado la API Key de Gemini. Ve a Ajustes para configurarla.")

        if not file_paths and not raw_text.strip():
            raise Exception("No has seleccionado archivos ni ingresado apuntes para compilar.")

        parts = []
        text_corpus = []

        # 1. Procesar cada archivo de la lista
        total_files = len(file_paths)
        for idx, fpath in enumerate(file_paths, start=1):
            fname = os.path.basename(fpath)
            if progress_callback:
                progress_callback(f"[{idx}/{total_files}] Leyendo archivo: {fname}...")

            extracted_text, media_mime, base64_data = self.extract_text_from_file(fpath)

            if base64_data and media_mime:
                parts.append({
                    "inline_data": {
                        "mime_type": media_mime,
                        "data": base64_data
                    }
                })
                parts.append({
                    "text": f"--- ADJUNTO MULTIMEDIA {idx}: {fname} ({media_mime}) ---"
                })
            else:
                text_corpus.append(f"=== FUENTE {idx} ({fname}) ===\n{extracted_text}\n")

        # 2. Agregar texto directo si existe
        if raw_text.strip():
            text_corpus.append(f"=== APUNTES DIRECTOS / NOTAS ADICIONALES ===\n{raw_text.strip()}\n")

        # 3. Consolidar prompt de usuario
        if text_corpus:
            consolidated_text = "\n".join(text_corpus)
            parts.insert(0, {
                "text": f"A CONTINUACIÓN SE PRESENTAN LOS APUNTES Y DOCUMENTOS RECOPILADOS DE LA REUNIÓN (COMPILACIÓN DE {len(file_paths)} FUENTES). Analízalos conjuntamente para consolidar una única minuta formal y completa:\n\n{consolidated_text}"
            })
        else:
            parts.insert(0, {
                "text": "Analiza los archivos adjuntos compilados de la reunión y genera la minuta estructurada según las reglas del sistema."
            })

        # 4. Enviar a Gemini con fallback inteligente
        models_to_try = [self.model_name]
        for m in (self.cached_models or ["gemini-2.0-flash", "gemini-1.5-flash", "gemini-1.5-pro"]):
            if m not in models_to_try:
                models_to_try.append(m)

        last_error = None
        raw_output = None

        for model in models_to_try:
            try:
                if progress_callback:
                    progress_callback(f"Sintetizando minuta con {model}...")
                raw_output = self._call_gemini_api(model, parts, system_prompt)
                if raw_output:
                    break
            except Exception as e:
                last_error = e
                if "404" in str(e) or "not found" in str(e).lower() or "unsupported" in str(e).lower():
                    continue
                else:
                    raise e

        if not raw_output:
            raise last_error or Exception("No se pudo obtener respuesta de los modelos de Gemini.")

        if progress_callback:
            progress_callback("Estructurando minuta y acuerdos extraídos...")

        clean_text = raw_output.strip()
        clean_text = re.sub(r"^```json\s*", "", clean_text)
        clean_text = re.sub(r"^```\s*", "", clean_text)
        clean_text = re.sub(r"\s*```$", "", clean_text)

        try:
            data = json.loads(clean_text)
        except Exception:
            match = re.search(r"\{.*\}", clean_text, re.DOTALL)
            if match:
                try:
                    data = json.loads(match.group(0))
                except Exception:
                    raise Exception("La IA no devolvió un formato JSON procesable.")
            else:
                raise Exception("La IA no devolvió una estructura JSON válida.")

        data.setdefault("fecha", "No especificado")
        data.setdefault("hora", "No especificado")
        data.setdefault("lugar", "No especificado")
        data.setdefault("asunto", "Minuta de Reunión")
        data.setdefault("coordinador", "Coordinador(a) de Reunión")
        data.setdefault("detalles", "Sin detalles registrados.")
        data.setdefault("asistentes", "No especificado")
        data.setdefault("filas", [])

        standard_filas = []
        for f in data.get("filas", []):
            if isinstance(f, dict):
                tema = str(f.get("tema", f.get("TEMA", "Punto tratado"))).strip()
                comp = str(f.get("compromiso", f.get("COMPROMISO", "Acuerdo"))).strip()
                resp = str(f.get("responsable", f.get("RESPONSABLE", f.get("plazo", "No especificado")))).strip()
                plazo = str(f.get("plazo", f.get("PLAZO", "No especificado"))).strip()
                
                standard_filas.append({
                    "tema": tema or "Punto tratado",
                    "compromiso": comp or "Acuerdo",
                    "responsable": resp or "No especificado",
                    "plazo": plazo or "No especificado"
                })
        data["filas"] = standard_filas

        return data
